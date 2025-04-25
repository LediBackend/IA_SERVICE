from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
import aiofiles
import os
import hashlib
import pdfplumber
import docx

async def processedAndSaved(uploadFolder,file,fileName):
     
     FilePath = os.path.join(uploadFolder,fileName)

     async with aiofiles.open(FilePath,'wb') as f:
          
          content = await file.read()
          
          await f.write(content)
     
     return FilePath

def getFileHash(filePath):
    
    hasher = hashlib.sha256()
    
    with open(filePath, "rb") as f:
    
        buf = f.read()
    
        hasher.update(buf)
    
    return hasher.hexdigest()

def isPdfAlreadyIndexed(fileHash,fVector):
    
    result = fVector.similarity_search(fileHash, k=1)
    
    if result:
    
        for doc in result:
    
            if doc.metadata.get("file_hash") == fileHash:
    
                return True
    
    return False

def loadFile(filePath):
     
     fileExt = filePath.split('.')[-1].lower()

     if fileExt == "txt":
          
          with open(filePath,'r',encoding='utf-8') as f:
          
               content =f.read()

     elif fileExt == 'pdf':
          
          with pdfplumber.open(filePath) as pdf:
               
               content ='\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
     
     elif fileExt == "docx":

          doc = docx.Document(filePath)

          content = '\n'.join([p.text for p in doc.paragraphs])

     else:

          raise ValueError('Archivo proporcionado no es posible procesar')
     
     return content

def textSplitter(document, bookName=None):

     textSplitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100, add_start_index=True )
     
     if isinstance(document, str):
        
        document = [Document(page_content=document)]

     chunks = textSplitter.split_documents(document)
     
     if bookName:
     
          for chunk in chunks:
     
               chunk.page_content = f'Curso:{bookName}\n{chunk.page_content}'

     return chunks

def indexDocs(documents, fVector):
    
    fVector.add_documents(documents)
